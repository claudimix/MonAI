import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from gtts import gTTS
st.set_page_config(
    page_title="Assistant IA YouTube",
    page_icon="🎬",
    layout="wide"
)

st.title("🎬 Mon Assistant IA avec Gemini")
st.write("Crée des Shorts, génère des prompts d’images, apprends le portugais, explore le trading.")

with st.sidebar:
    st.header("⚙️ Paramètres")
    api_key = st.text_input("Clé Gemini", type="password")

    mode = st.selectbox(
        "Choisissez un mode",
        ["Général", "Portugais", "Trading", "YouTube"]
    )

    st.info("Ne partage jamais ta clé Gemini publiquement.")
    st.write("Créé par Jean-Claude")

question = st.text_area("✍️ Pose-moi une question ou écris le sujet de ton Short")

if "historique" not in st.session_state:
    st.session_state.historique = []

if "dernier_short" not in st.session_state:
    st.session_state.dernier_short = ""

if "derniers_prompts_images" not in st.session_state:
    st.session_state.derniers_prompts_images = ""


def obtenir_role(mode):
    if mode == "YouTube":
        return """
Tu es un expert YouTube Shorts.
Tu aides à créer des titres, hooks, scripts, scènes, prompts d'images,
descriptions, hashtags, miniatures, mots-clés SEO et conseils CapCut.
"""
    elif mode == "Trading":
        return """
Tu es un expert trading.
Explique clairement le Forex, les cryptomonnaies, les actions et la gestion du risque.
Rappelle toujours que le trading comporte des risques.
"""
    elif mode == "Portugais":
        return """
Tu es un professeur de portugais brésilien pour francophones.
Tu aides à traduire, corriger, expliquer la grammaire et enseigner la conjugaison.
"""
    else:
        return """
Tu es un assistant IA général.
Tu réponds clairement en français.
"""


def demander_gemini(api_key, mode, question):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-flash-latest")

    ancien_contexte = ""
    for item in st.session_state.historique:
        ancien_contexte += f"Utilisateur : {item['question']}\n"
        ancien_contexte += f"Assistant : {item['reponse']}\n\n"

    prompt = f"""
{obtenir_role(mode)}

Historique :
{ancien_contexte}

Nouvelle question :
{question}

Réponds en français clair, utile et professionnel.
"""

    response = model.generate_content(prompt)
    return response.text


def creer_document_word(texte):
    document = Document()
    document.add_heading("Document généré par Mon Assistant IA", level=1)
    document.add_paragraph("Créé avec l'assistant IA de Jean-Claude")
    document.add_paragraph("")

    for ligne in texte.split("\n"):
        ligne = ligne.strip()

        if not ligne:
            continue

        ligne = ligne.replace("###", "")
        ligne = ligne.replace("##", "")
        ligne = ligne.replace("#", "")
        ligne = ligne.replace("**", "")
        ligne = ligne.replace("---", "")
        ligne = ligne.replace(">", "")

        if ligne.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "11.", "12.")):
            document.add_heading(ligne, level=2)
        else:
            document.add_paragraph(ligne)

    fichier = BytesIO()
    document.save(fichier)
    fichier.seek(0)
    return fichier


col1, col2, col3, col4 = st.columns(4)

with col1:
    bouton_repondre = st.button("🤖 Répondre")

with col2:
    bouton_short = st.button("🎬 Créer un Short complet")

with col3:
    bouton_images = st.button("🎨 Prompts d’images IA")

with col4:
    bouton_effacer = st.button("🗑️ Effacer l’historique")


if bouton_repondre:
    if not api_key:
        st.warning("Veuillez coller votre clé Gemini dans la barre latérale.")
    elif not question:
        st.warning("Veuillez écrire une question.")
    else:
        reponse = demander_gemini(api_key, mode, question)

        st.session_state.historique.append({
            "question": question,
            "reponse": reponse
        })


if bouton_short:
    if not api_key:
        st.warning("Veuillez coller votre clé Gemini dans la barre latérale.")
    elif not question:
        st.warning("Écris d'abord le sujet du Short.")
    else:
        prompt_short = f"""
Crée un YouTube Short complet sur ce sujet :

{question}

Donne exactement cette structure :

1. Idée générale de la vidéo

2. Trois titres YouTube accrocheurs :
- Titre 1 :
- Titre 2 :
- Titre 3 :

3. Titre court pour la miniature :
Un titre très court, puissant, facile à lire sur une miniature.

4. Hook de 3 secondes :
Une phrase très forte pour retenir l’attention immédiatement.

5. Script complet de 30 à 45 secondes :
Écris un script fluide, naturel, adapté à une voix off.

6. Découpage en 5 scènes :
Pour chaque scène, donne :
- durée approximative ;
- action visuelle ;
- texte à afficher à l’écran.

7. Prompts d’images IA pour chaque scène :
Pour chaque scène, donne un prompt détaillé avec :
- sujet principal ;
- décor ;
- style visuel ;
- lumière ;
- ambiance ;
- format vertical 9:16 ;
- haute résolution ;
- aucun texte dans l’image.

8. Description YouTube optimisée :
Écris une description claire, attirante et optimisée pour le référencement.

9. Quinze hashtags :
Donne exactement 15 hashtags pertinents.

10. Cinq mots-clés SEO :
Donne exactement 5 mots-clés ou expressions recherchables.

11. Idée de miniature :
Explique clairement l’image de miniature idéale.

12. Conseil de montage CapCut :
Donne des conseils simples pour le rythme, les transitions, les sous-titres et la musique.

Réponds en français clair, professionnel et directement exploitable.
"""

        reponse = demander_gemini(api_key, "YouTube", prompt_short)
        st.session_state.dernier_short = reponse

        st.session_state.historique.append({
            "question": "Créer un Short YouTube complet : " + question,
            "reponse": reponse
        })


if bouton_images:
    if not api_key:
        st.warning("Veuillez coller votre clé Gemini dans la barre latérale.")
    elif not question:
        st.warning("Écris d'abord le sujet ou le Short dans la zone de question.")
    else:
        prompt_images = f"""
Crée 5 prompts d'images IA pour un YouTube Short sur ce sujet :

{question}

Pour chaque prompt, donne exactement cette structure :

Image 1 :
Prompt :
Texte conseillé à l'écran :

Image 2 :
Prompt :
Texte conseillé à l'écran :

Image 3 :
Prompt :
Texte conseillé à l'écran :

Image 4 :
Prompt :
Texte conseillé à l'écran :

Image 5 :
Prompt :
Texte conseillé à l'écran :

Chaque prompt doit être :
- très détaillé ;
- en format vertical 9:16 ;
- haute résolution ;
- style cinématographique ;
- adapté à YouTube Shorts ;
- sans texte dans l'image ;
- utilisable dans ChatGPT, Gemini, Canva, Leonardo ou autre IA d'image.

Réponds en français.
"""

        reponse = demander_gemini(api_key, "YouTube", prompt_images)
        st.session_state.derniers_prompts_images = reponse

        st.session_state.historique.append({
            "question": "Prompts d’images IA : " + question,
            "reponse": reponse
        })


if bouton_effacer:
    st.session_state.historique = []
    st.session_state.dernier_short = ""
    st.session_state.derniers_prompts_images = ""


st.subheader("💬 Historique de conversation")

for item in st.session_state.historique:
    st.markdown("### Vous")
    st.write(item["question"])

    st.markdown("### Assistant")
    st.write(item["reponse"])

    st.divider()


if st.session_state.dernier_short:
    st.subheader("📥 Télécharger le dernier Short")

    st.download_button(
        label="💾 Télécharger le Short en TXT",
        data=st.session_state.dernier_short,
        file_name="mon_short_youtube.txt",
        mime="text/plain"
    )

    fichier_word = creer_document_word(st.session_state.dernier_short)

    st.download_button(
        label="📄 Télécharger le Short en Word",
        data=fichier_word,
        file_name="mon_short_youtube.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if st.session_state.derniers_prompts_images:
    st.subheader("📥 Télécharger les prompts d’images")

    st.download_button(
        label="🎨 Télécharger les prompts d’images en TXT",
        data=st.session_state.derniers_prompts_images,
        file_name="prompts_images_youtube.txt",
        mime="text/plain"
    )

    fichier_word_images = creer_document_word(st.session_state.derniers_prompts_images)

    st.download_button(
        label="📄 Télécharger les prompts d’images en Word",
        data=fichier_word_images,
        file_name="prompts_images_youtube.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if st.session_state.dernier_short:

    tts = gTTS(text=st.session_state.dernier_short, lang="fr")

    audio_buffer = BytesIO()
    tts.write_to_fp(audio_buffer)
    audio_buffer.seek(0)

    st.download_button(
        label="🎤 Télécharger le Short en MP3",
        data=audio_buffer,
        file_name="mon_short_youtube.mp3",
        mime="audio/mpeg"
    )